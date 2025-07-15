"""
HOTSOFT火热口算题生成器 - 主程序
"""

import gui
import win32ui
import question
import docx
from docx.shared import Pt
from docx.oxml.shared import qn
from docx.oxml import OxmlElement


def main():
    # 获取用户设置
    user_settings = gui.userlist()
    if not user_settings:
        print("用户取消了操作")
        return

    print("用户设置:", user_settings)

    # 生成题目
    questions = []
    if user_settings[0] == 1:  # 整数
        questions = question.question(
            int(user_settings[2]),
            int(user_settings[3]),
            user_settings[1],
            1
        )
    elif user_settings[0] == 2:  # 小数
        questions = question.question(
            float(user_settings[2]),
            float(user_settings[3]),
            user_settings[1],
            2
        )
    elif user_settings[0] == 3:  # 分数
        questions = question.question(
            int(user_settings[2]),
            int(user_settings[3]),
            user_settings[1],
            3
        )

    print("生成的题目:", questions)

    if not questions:
        print("错误：没有生成任何题目")
        return

    # 选择保存路径
    dlg = win32ui.CreateFileDialog(0)
    dlg.SetOFNTitle("把题目保存在哪里呢？")
    if dlg.DoModal() != 1:  # 1表示点击了确定
        print("用户取消了文件选择")
        return

    file_path = dlg.GetPathName()
    print("保存路径:", file_path)

    if not file_path:
        print("错误：没有选择保存路径")
        return

    # 创建Word文档
    doc = docx.Document()

    # 设置两栏布局
    section = doc.sections[0]
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')
    sectPr.append(cols)

    # 添加题目到文档（每道题后加5个空段落）
    for i, q in enumerate(questions, 1):
        # 添加题目
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(f"{i}. {q}")
        run.font.size = Pt(12)
        run.font.name = '宋体'

        # 添加5个空段落
        for _ in range(5):
            doc.add_paragraph()

    # 保存文档
    try:
        doc.save(file_path)
        print(f"文档已成功保存到: {file_path}")
    except Exception as e:
        print(f"保存文档时出错: {e}")


if __name__ == "__main__":
    main()